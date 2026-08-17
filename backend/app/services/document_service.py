import io
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib import colors

from app.schemas.resume import StructuredResume


class DocumentGenerationService:

    # -------------------------------------------------------------------------
    # DOCX: Resume Generator
    # -------------------------------------------------------------------------
    @staticmethod
    def generate_resume_docx(resume: StructuredResume) -> io.BytesIO:
        doc = Document()

        for s in doc.sections:
            s.top_margin = Inches(0.75)
            s.bottom_margin = Inches(0.75)
            s.left_margin = Inches(0.75)
            s.right_margin = Inches(0.75)

        # Header: Name & Contact
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(resume.contact_info.full_name)
        name_run.bold = True
        name_run.font.size = Pt(18)
        name_run.font.name = "Calibri"

        contact_items = []
        if resume.contact_info.email:
            contact_items.append(str(resume.contact_info.email))
        if resume.contact_info.phone:
            contact_items.append(resume.contact_info.phone)
        if resume.contact_info.location:
            contact_items.append(resume.contact_info.location)
        if resume.contact_info.linkedin_url:
            contact_items.append(str(resume.contact_info.linkedin_url))
        if resume.contact_info.github_url:
            contact_items.append(str(resume.contact_info.github_url))

        if contact_items:
            c_para = doc.add_paragraph(" | ".join(contact_items))
            c_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if c_para.runs:
                c_para.runs[0].font.size = Pt(10)
                c_para.runs[0].font.name = "Calibri"
            c_para.paragraph_format.space_after = Pt(10)

        def add_section_header(title: str):
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(2)
            run = h.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(30, 41, 59)

        # Summary
        if resume.professional_summary:
            add_section_header("Professional Summary")
            p = doc.add_paragraph(resume.professional_summary)
            if p.runs:
                p.runs[0].font.size = Pt(10)
                p.runs[0].font.name = "Calibri"
            p.paragraph_format.space_after = Pt(6)

        # Experience
        if resume.work_experience:
            add_section_header("Work Experience")
            for exp in resume.work_experience:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)

                r_title = p.add_run(f"{exp.job_title} ")
                r_title.bold = True
                r_title.font.size = Pt(10.5)
                r_title.font.name = "Calibri"

                r_comp = p.add_run(f"- {exp.company_name}")
                r_comp.font.size = Pt(10)
                r_comp.font.name = "Calibri"

                date_loc = []
                if exp.start_date or exp.end_date:
                    date_loc.append(f"{exp.start_date or ''} - {exp.end_date or 'Present'}")
                if exp.location:
                    date_loc.append(exp.location)

                if date_loc:
                    p2 = doc.add_paragraph(" | ".join(date_loc))
                    if p2.runs:
                        p2.runs[0].italic = True
                        p2.runs[0].font.size = Pt(9.5)
                        p2.runs[0].font.name = "Calibri"
                    p2.paragraph_format.space_after = Pt(2)

                for bp in exp.bullet_points:
                    bp_para = doc.add_paragraph(bp, style="List Bullet")
                    bp_para.paragraph_format.space_after = Pt(1)
                    bp_para.paragraph_format.left_indent = Inches(0.25)
                    if bp_para.runs:
                        bp_para.runs[0].font.size = Pt(10)
                        bp_para.runs[0].font.name = "Calibri"

        # Education
        if resume.education:
            add_section_header("Education")
            for edu in resume.education:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(1)

                r_deg = p.add_run(f"{edu.degree} ")
                r_deg.bold = True
                r_deg.font.size = Pt(10.5)
                r_deg.font.name = "Calibri"

                r_inst = p.add_run(f"- {edu.institution}")
                r_inst.font.size = Pt(10)
                r_inst.font.name = "Calibri"

                if edu.graduation_year:
                    p_year = doc.add_paragraph(f"Graduated: {edu.graduation_year}")
                    if p_year.runs:
                        p_year.runs[0].italic = True
                        p_year.runs[0].font.size = Pt(9.5)
                        p_year.runs[0].font.name = "Calibri"

        # Skills
        if resume.skills:
            add_section_header("Skills")
            sk = resume.skills
            if sk.technical_skills:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("Technical Skills: ")
                r.bold = True
                r.font.size = Pt(10)
                p.add_run(", ".join(sk.technical_skills)).font.size = Pt(10)
            if sk.tools_and_frameworks:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("Tools & Frameworks: ")
                r.bold = True
                r.font.size = Pt(10)
                p.add_run(", ".join(sk.tools_and_frameworks)).font.size = Pt(10)
            if sk.soft_skills:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                r = p.add_run("Soft Skills: ")
                r.bold = True
                r.font.size = Pt(10)
                p.add_run(", ".join(sk.soft_skills)).font.size = Pt(10)

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # -------------------------------------------------------------------------
    # DOCX: Cover Letter Generator
    # -------------------------------------------------------------------------
    @staticmethod
    def generate_cover_letter_docx(
        candidate_name: str,
        contact_info_line: str,
        company_name: str,
        content: str
    ) -> io.BytesIO:
        doc = Document()
        for s in doc.sections:
            s.top_margin = Inches(1.0)
            s.bottom_margin = Inches(1.0)
            s.left_margin = Inches(1.0)
            s.right_margin = Inches(1.0)

        # Candidate Header
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(candidate_name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = "Calibri"

        if contact_info_line:
            c_p = doc.add_paragraph(contact_info_line)
            if c_p.runs:
                c_p.runs[0].font.size = Pt(10)
                c_p.runs[0].font.name = "Calibri"
            c_p.paragraph_format.space_after = Pt(14)

        # Body paragraphs
        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        for text_p in paragraphs:
            p = doc.add_paragraph(text_p)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            if p.runs:
                p.runs[0].font.size = Pt(11)
                p.runs[0].font.name = "Calibri"

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # -------------------------------------------------------------------------
    # PDF: Resume Generator
    # -------------------------------------------------------------------------
    @staticmethod
    def generate_resume_pdf(resume: StructuredResume) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.6 * inch,
            rightMargin=0.6 * inch,
            topMargin=0.6 * inch,
            bottomMargin=0.6 * inch,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "ResumeTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            alignment=1,
            textColor=colors.HexColor("#0f172a"),
        )
        contact_style = ParagraphStyle(
            "ResumeContact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            alignment=1,
            textColor=colors.HexColor("#475569"),
        )
        sec_heading = ParagraphStyle(
            "SectionHeading",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=12,
            spaceBefore=6,
            spaceAfter=2,
            textColor=colors.HexColor("#1e293b"),
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#1e293b"),
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            leading=11.5,
            leftIndent=12,
            firstLineIndent=-8,
            textColor=colors.HexColor("#1e293b"),
        )

        story = []

        # Name & Contact Header
        story.append(Paragraph(resume.contact_info.full_name, title_style))
        contact_items = []
        if resume.contact_info.email:
            contact_items.append(str(resume.contact_info.email))
        if resume.contact_info.phone:
            contact_items.append(resume.contact_info.phone)
        if resume.contact_info.location:
            contact_items.append(resume.contact_info.location)
        if resume.contact_info.linkedin_url:
            contact_items.append(str(resume.contact_info.linkedin_url))
        if resume.contact_info.github_url:
            contact_items.append(str(resume.contact_info.github_url))

        if contact_items:
            story.append(Spacer(1, 2))
            story.append(Paragraph(" &bull; ".join(contact_items), contact_style))

        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#cbd5e1"), spaceAfter=4))

        # Summary
        if resume.professional_summary:
            story.append(Paragraph("PROFESSIONAL SUMMARY", sec_heading))
            story.append(Paragraph(resume.professional_summary, body_style))
            story.append(Spacer(1, 4))

        # Experience
        if resume.work_experience:
            story.append(Paragraph("WORK EXPERIENCE", sec_heading))
            for exp in resume.work_experience:
                date_str = f"{exp.start_date or ''} - {exp.end_date or 'Present'}"
                job_line = f"<b>{exp.job_title}</b> - {exp.company_name} <font color='#64748b'>({date_str})</font>"
                story.append(Paragraph(job_line, body_style))
                for bp in exp.bullet_points:
                    story.append(Paragraph(f"&bull; {bp}", bullet_style))
                story.append(Spacer(1, 3))

        # Education
        if resume.education:
            story.append(Paragraph("EDUCATION", sec_heading))
            for edu in resume.education:
                grad = f" ({edu.graduation_year})" if edu.graduation_year else ""
                edu_line = f"<b>{edu.degree}</b> - {edu.institution}{grad}"
                story.append(Paragraph(edu_line, body_style))
            story.append(Spacer(1, 3))

        # Skills
        if resume.skills:
            story.append(Paragraph("SKILLS", sec_heading))
            sk = resume.skills
            if sk.technical_skills:
                story.append(Paragraph(f"<b>Technical Skills:</b> {', '.join(sk.technical_skills)}", body_style))
            if sk.tools_and_frameworks:
                story.append(Paragraph(f"<b>Tools & Frameworks:</b> {', '.join(sk.tools_and_frameworks)}", body_style))
            if sk.soft_skills:
                story.append(Paragraph(f"<b>Soft Skills:</b> {', '.join(sk.soft_skills)}", body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer

    # -------------------------------------------------------------------------
    # PDF: Cover Letter Generator
    # -------------------------------------------------------------------------
    @staticmethod
    def generate_cover_letter_pdf(
        candidate_name: str,
        contact_info_line: str,
        company_name: str,
        content: str
    ) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            leftMargin=0.8 * inch,
            rightMargin=0.8 * inch,
            topMargin=0.8 * inch,
            bottomMargin=0.8 * inch,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "CLTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=16,
            leading=18,
            textColor=colors.HexColor("#0f172a"),
        )
        contact_style = ParagraphStyle(
            "CLContact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            leading=13,
            textColor=colors.HexColor("#475569"),
        )
        body_style = ParagraphStyle(
            "CLBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=10,
            leading=14,
            spaceAfter=8,
            textColor=colors.HexColor("#1e293b"),
        )

        story = []
        story.append(Paragraph(candidate_name, title_style))
        if contact_info_line:
            story.append(Paragraph(contact_info_line, contact_style))

        story.append(Spacer(1, 10))
        story.append(HRFlowable(width="100%", thickness=0.75, color=colors.HexColor("#cbd5e1"), spaceAfter=14))

        paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
        for p in paragraphs:
            story.append(Paragraph(p.replace("\n", "<br/>"), body_style))

        doc.build(story)
        buffer.seek(0)
        return buffer
