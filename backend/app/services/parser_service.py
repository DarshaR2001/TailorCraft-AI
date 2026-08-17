import io
import pdfplumber
import pytesseract
from docx import Document
from pdf2image import convert_from_bytes


class DocumentParsingService:
    @staticmethod
    def parse_pdf(file_bytes: bytes) -> str:
        extracted_text = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    extracted_text.append(text)
        full_text = "\n".join(extracted_text).strip()
        if len(full_text) < 50:
            full_text = DocumentParsingService._ocr_pdf(file_bytes)
        return full_text

    @staticmethod
    def _ocr_pdf(file_bytes: bytes) -> str:
        images = convert_from_bytes(file_bytes)
        ocr_text = [pytesseract.image_to_string(img) for img in images]
        return "\n".join(ocr_text).strip()

    @staticmethod
    def parse_docx(file_bytes: bytes) -> str:
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs).strip()

    @staticmethod
    def parse_txt(file_bytes: bytes) -> str:
        return file_bytes.decode("utf-8", errors="ignore").strip()

    @classmethod
    def extract_text(cls, filename: str, file_bytes: bytes) -> str:
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            return cls.parse_pdf(file_bytes)
        elif lower_name.endswith(".docx"):
            return cls.parse_docx(file_bytes)
        elif lower_name.endswith(".txt"):
            return cls.parse_txt(file_bytes)
        else:
            raise ValueError("Unsupported file format. Accepted: PDF, DOCX, TXT.")
